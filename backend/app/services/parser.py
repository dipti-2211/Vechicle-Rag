"""
Vehicle Intelligence Assistant — Document Parser Service

Responsible for extracting raw text from uploaded files (PDF, CSV, XLSX, TXT, DOCX).
"""

import logging
from pathlib import Path

import pandas as pd
from pypdf import PdfReader

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service to parse and extract text from various document formats."""

    @staticmethod
    def parse(file_path: str, file_type: str) -> str:
        """
        Route to the correct parser based on file_type.

        Args:
            file_path: Absolute or relative path to the file.
            file_type: Extension (pdf, csv, xlsx, txt).

        Returns:
            Extracted text as a single string.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower()
        
        try:
            if file_type == 'pdf':
                return DocumentParser._parse_pdf(path)
            elif file_type == 'csv':
                return DocumentParser._parse_csv(path)
            elif file_type == 'xlsx':
                return DocumentParser._parse_excel(path)
            elif file_type == 'txt':
                return DocumentParser._parse_txt(path)
            elif file_type == 'docx':
                return DocumentParser._parse_docx(path)
            else:
                raise ValueError(f"Unsupported file type for parsing: {file_type}")
        except Exception as e:
            logger.error("Error parsing %s: %s", file_path, e)
            raise RuntimeError(f"Failed to parse document: {e}")

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Extract text from a PDF file."""
        text = []
        with open(path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n\n".join(text)

    @staticmethod
    def _parse_csv(path: Path) -> str:
        """Extract text from a CSV by converting rows to readable text."""
        for encoding in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
            try:
                df = pd.read_csv(path, encoding=encoding)
                # Drop unnamed columns (from merged cells / empty headers)
                df = df.loc[:, ~df.columns.astype(str).str.match(r'^Unnamed')]
                return DocumentParser._df_to_text(df)
            except (UnicodeDecodeError, Exception):
                continue
        raise RuntimeError(f"Could not read CSV file with any supported encoding: {path}")

    @staticmethod
    def _parse_excel(path: Path) -> str:
        """
        Extract text from an Excel file (all sheets concatenated).

        Uses header=None to avoid the 'Unnamed: N' column problem caused by
        merged title cells. Converts each non-empty row into a readable pipe-
        separated string so Gemini can understand tabular data.
        """
        try:
            # header=None: treat all rows as data, columns are 0, 1, 2...
            all_sheets = pd.read_excel(path, sheet_name=None, header=None)
            parts = []
            for sheet_name, df in all_sheets.items():
                parts.append(f"=== Sheet: {sheet_name} ===")
                parts.append(DocumentParser._df_noheader_to_text(df))
            result = "\n\n".join(parts)
            if result.strip():
                return result
        except Exception as e:
            logger.warning("Excel header=None parsing failed: %s — trying default", e)

        # Fallback: standard header parsing (drops Unnamed columns)
        try:
            all_sheets = pd.read_excel(path, sheet_name=None)
            parts = []
            for sheet_name, df in all_sheets.items():
                df = df.loc[:, ~df.columns.astype(str).str.match(r'^Unnamed')]
                parts.append(f"=== Sheet: {sheet_name} ===")
                parts.append(DocumentParser._df_to_text(df))
            return "\n\n".join(parts)
        except Exception:
            df = pd.read_excel(path)
            return df.to_string(index=False)

    @staticmethod
    def _df_noheader_to_text(df: pd.DataFrame) -> str:
        """Convert a header=None DataFrame to pipe-separated readable lines."""
        lines = []
        for _, row in df.iterrows():
            cells = [str(v).strip() for v in row if pd.notna(v) and str(v).strip() not in ('', 'nan', 'NaN')]
            if cells:
                lines.append(" | ".join(cells))
        return "\n".join(lines)

    @staticmethod
    def _df_to_text(df: pd.DataFrame) -> str:
        """Convert a DataFrame with named columns to readable 'Col: Value' lines."""
        cols = df.columns.tolist()
        lines = []
        for _, row in df.iterrows():
            parts = []
            for col in cols:
                val = row[col]
                if pd.notna(val) and str(val).strip() not in ('', 'nan', 'NaN'):
                    parts.append(f"{col}: {val}")
            if parts:
                lines.append(", ".join(parts))
        return "\n".join(lines)

    @staticmethod
    def _parse_txt(path: Path) -> str:
        """Read plain text, trying common encodings."""
        for encoding in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        # Last resort: read as bytes and decode with errors='replace'
        with open(path, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Read DOCX files including paragraphs and tables."""
        try:
            import docx
            doc = docx.Document(path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    text = [cell.text for cell in row.cells]
                    content.append(" | ".join(text))
            return "\n".join(content)
        except ImportError:
            logger.warning("python-docx not installed. DOCX parsing will fail.")
            raise RuntimeError("python-docx is required to parse .docx files")
